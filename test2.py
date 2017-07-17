import docx
import os
# test file
docx_name = input('Please input File Name ：')
docx_name = docx_name+ '.docx'
print(docx_name)
document = docx.Document('example2.docx')

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'

table = document.tables

def auto_word_cell(replaced_word,content):
    for row in table[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replaced_word in paragraph.text:
                        temp = paragraph.text
                        print(temp)
                        temp = temp.replace(replaced_word,content)
                        print(temp)
                        paragraph.text = temp
                        paragraph.style = 'Normal'
                        # print(inline[i].text)
                        # inline[i].bold = True

# row = table[0].rows[0]

#
# cell = row.cells[6]
#
# paragraph = cell.paragraphs[0]
# print(paragraph.text)
# style = document.styles['Normal']
# font = style.font
# font.name = 'Arial'
# table[0].style = document.styles['Normal']
for paragraph in document.paragraphs:
    inline = paragraph.runs
    for i in range(len(inline)):
        print(inline[i].text)
    # if 'name' in paragraph.text:
    #     print(paragraph.text)
    #     paragraph.text = 'new text containing ocean'

auto_word_cell('v_i','hahaha\n hahaha')
document.save(docx_name)
