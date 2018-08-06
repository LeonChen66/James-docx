# -*- coding: UTF-8 -*-
from PyQt5 import QtWidgets, QtGui
import sys
import docx
from os import mkdir,chdir,system
from james import Ui_Form    # import qt designer's gui
from datetime import datetime
from dateutil.relativedelta import relativedelta

class mywindow(QtWidgets.QWidget,Ui_Form):
    def __init__(self):
        super(mywindow,self).__init__()
        self.setupUi(self)
        self.document = docx.Document('example2.docx')
        self.table = self.document.tables
    # click save function
    def hello(self):
        self.document = docx.Document('example2.docx')
        self.table = self.document.tables
        self.nessesary_information()
        self.basic_information()
        self.mitral()
        self.aortic()
        self.tricuspid()
        self.pulmonary()
        self.wall_motion()
        self.comment_word()
        self.save_docx()
        # self.document.save('test_james.docx')
        # os.system('test_james.docx')

    def save_docx(self):
        dir_name = self.ID_input.text()
        file_name = self.ID_input.text()+datetime.now().strftime('_%Y_%m_%d') + '.docx'
        try:
            mkdir(dir_name)
            chdir(dir_name)
            self.document.save(file_name)
            system(file_name)
            chdir("..")
        except:
            chdir(dir_name)
            self.document.save(file_name)
            system(file_name)
            chdir("..")

    def mitral(self):
        content = ''
        if self.M_normal_check.isChecked():
            content += '> ' + self.M_normal_check.text() + '\n'

        if self.M_BMV_check.isChecked():
            content += '> ' + self.M_BMV_check.text() + '\n'

        if self.M_MMV_check.isChecked():
            content += '> '+self.M_MMV_check.text()+'\n'

        if self.M_MMVp_check.isChecked():
            content += '> '+self.M_MMVp_check.text()+'\n'

        if self.M_MMVp_combo1.currentIndex() !=0:
            content += ': Anterior leaflet ==>' + self.M_MMVp_combo1.currentText()

        if self.M_MMVp_combo.currentIndex() !=0:
            content += ': Posterior leaflet  ==>' + self.M_MMVp_combo.currentText()+'\n'

        if self.M_mitral_check.isChecked():
            content += '> Mitral annulus calcification\n'

        if self.M_scler_check.isChecked():
            content += '> '+self.M_scler_check.text()+'\n'

        if self.M__fibro_check.isChecked():
            content += '> '+self.M__fibro_check.text()+'\n'

        if self.M_myxo_check.isChecked():
            content += '> '+self.M_myxo_check.text()+'\n'

        if self.M_scler_check.isChecked():
            content += '> '+self.M_scler_check.text()+'\n'

        if self.M_chordae_check.isChecked():
            content += '> '+self.M_chordae_check.text()+'  '+\
                       self.M_chordae_check_2.currentText()+'\n'

        if self.M_MVA_check.isChecked():
            content += '> '+self.M_MVA_check.text()+' : '+\
                self.M_MVA_input.text()+' cm^2\n'

        if self.M_meantrans_check.isChecked():
            content += '> '+self.M_meantrans_check.text()+' MV P.G.: '+\
                self.M_meantrans_input.text()+' mmHg\n'

        if self.M_MR_check.isChecked():
            content += '  '+self.M_MR_check.text()+' : '+\
                self.M_MR_combo.currentText()+''
            content += '   >> MR area : ' + self.M_MRarea_input.text() + ' cm^2'
            content += '   >> MR area/LA area : ' + self.M_MRLA_input.text() + ' %'
            content += '   >> Veno contrata: ' + self.M_Veno_input.text() + ' mm'
            content += '   >> ERO : ' + self.M_ERO_input.text() + ' cm^2\n'

        if self.M_others_check.isChecked():
            content +='> Others : ' + self.M_others.toPlainText()

        self.auto_paragraph('MM',content)

    def aortic(self):
        content = ''
        if self.A_normal_check.isChecked():
            content += '> '+self.A_normal_check.text()+'\n'

        if self.A_BAV_check.isChecked():
            content += '> ' + 'Bioprothetic AV' + '\n'

        if self.A_MAV_check.isChecked():
            content += '> '+ self.A_MAV_check.text() +'\n'

        if self.A_scler_check.isChecked():
            content += '> '+self.A_scler_check.text()+'\n'

        if self.A_bicu_check.isChecked():
            content += '> '+self.A_bicu_check.text()+'\n'

        if self.A_tricu_check.isChecked():
            content += '> '+self.A_tricu_check.text()+'\n'

        if self.A_AVA_check.isChecked():
            content += '> '+self.A_AVA_check.text()+' : '+\
                self.A_AVA_input.text()+' cm^2\n'

        if self.A_meantrans_check.isChecked():
            content += '   >>'+self.A_meantrans_check.text()
            content += ' AV PG : '+self.A_AVPG_input.text()+' mmHg'
            content += '   >> peak trans AV PG : ' + self.A_peckPG_input.text() + ' mmHg'
            content += '   >> AV Vmax : ' + self.A_AVVmax_input.text() + ' m/s\n'

        if self.A_AR_check.isChecked():
            content += '> AR : ('+self.A_AR_combo.currentText() +')'
            content += '  >> jet height ratio = ' + self.A_jet_input.text()+' %'
            content += '  >> PHT = ' + self.A_PHT_input.text()+' ms'
            content += '  >> Veno contrata = ' + self.A_veno_input.text()+' mm'
            content += '  >> '+self.A_dias_combo.currentText()+' diastolic reversal flow\n'

        if self.A_others_check.isChecked():
            content += '> Others : ' + self.A_others.toPlainText()

        self.auto_paragraph('AA', content)


    def tricuspid(self):
        content = ''
        if self.T_normal_check.isChecked():
            content += '> Normal'

        if self.T_pro_check.isChecked():
            content += '\n> Prolapse : '

        if self.T_anter_check.isChecked():
            content += 'anterior leaflet, '

        if self.T_post_check.isChecked():
            content += 'posterior leaflet, '

        if self.T_septal_check.isChecked():
            content += 'septal leaflet '

        if self.T_TR_check.isChecked():
            content += '\n> TR : ' + self.T_TR_combo.currentText()
            content += '\n  >> TR area : '+self.T_TRarea_input.text()+ 'cm^2'

        if self.T_TRV_check.isChecked():
            content += '  > TRV : ' + self.T_TRV_input.text() + ' m/s'

        if self.T_TV_check.isChecked():
            content += '  > Trans-TV PG : ' + self.T_Trans_TV_input.text() + ' mm HG '

        if self.T_others_check.isChecked():
            content += '\n> Others : ' + self.T_others.toPlainText()


        content += '\n'
        self.auto_paragraph('TT',content)

    def pulmonary(self):
        content = ''
        if self.P_normal_check.isChecked():
            content += '> '+self.P_normal_check.text()

        if self.P_PR_check.isChecked():
            content += '\n> PR : ' + self.P_PR_combo.currentText()

        if self.P_others_check.isChecked():
            content += '\n> Others : ' + self.P_others.toPlainText()

        self.auto_paragraph('PP',content)

    def wall_motion(self):
        content = ''
        if self.W_normal_check.isChecked():
            content += self.W_normal_check.text()

        if self.W_abnormal_check.isChecked():
            content += self.W_abnormal_check.text() + ':'

        # -------- basal ---------
        content_1=''

        if self.W_basal_check.isChecked():
            content_1 += '' + self.W_basal_check.text() + ' ('

        if self.W_anterior_check.isChecked():
            content_1 += self.W_anterior_check.text() + ', '

        if self.W_septal_check.isChecked():
            content_1 += self.W_septal_check.text() + ', '

        if self.W_inferior_check.isChecked():
            content_1 +=self.W_inferior_check.text()+ ', '

        if self.W_posterior_check.isChecked():
            content_1 += self.W_posterior_check.text() + ', '

        if self.W_lateral_check.isChecked():
            content_1 += self.W_lateral_check.text() + ', '

        if self.W_abnormal_check.isChecked() and self.W_basal_check.isChecked():
            content_1 += ') '

        if self.W_hypo_check.isChecked():
            content_1 += self.W_hypo_check.text() + ', '

        if self.W_akinesis_check.isChecked():
            content_1 += self.W_akinesis_check.text() + ', '

        if self.W_dysk_check.isChecked():
            content_1 += self.W_dysk_check.text()



        # ------- midcavity ------
        content_2 = ''
        if self.W_midcavity_check.isChecked():
            content_2 += '' + self.W_midcavity_check.text() + ' ('

        if self.W_m_anterior_check.isChecked():
            content_2 += self.W_anterior_check.text() + ', '

        if self.W_m_septal_check.isChecked():
            content_2 += self.W_septal_check.text() + ', '

        if self.W_m_inferior_check.isChecked():
            content_2 += self.W_inferior_check.text()+ ', '

        if self.W_m_posterior_check.isChecked():
            content_2 += self.W_posterior_check.text() + ', '

        if self.W_m_lateral_check.isChecked():
            content_2 += self.W_lateral_check.text() + ', '

        if self.W_abnormal_check.isChecked() and self.W_midcavity_check.isChecked():
            content_2 += ') '

        if self.W_m_hypokinesis_check.isChecked():
            content_2 += self.W_hypo_check.text() + ', '

        if self.W_m_akinesis_check.isChecked():
            content_2 += self.W_akinesis_check.text() + ', '

        if self.W_m_dyskinesis_check.isChecked():
            content_2 += self.W_dysk_check.text()


        # ------- apical ----------
        content_3 = ''
        if self.W_apical_check.isChecked():
            content_3 += '' + self.W_apical_check.text() + ' ('
        if self.W_a_anterior_check.isChecked():
            content_3 += self.W_anterior_check.text() + ', '

        if self.W_a_septal_check.isChecked():
            content_3 += self.W_septal_check.text() + ', '

        if self.W_a_inferior_check.isChecked():
            content_3 +=self.W_inferior_check.text()+ ', '

        if self.W_a_posterior_check.isChecked():
            content_3 += self.W_posterior_check.text() + ', '

        if self.W_a_lateral_check.isChecked():
            content_3 += self.W_lateral_check.text() + ', '

        if self.W_abnormal_check.isChecked() and self.W_apical_check.isChecked():
            content_3 += ') '

        if self.W_a_hypokinesis_check.isChecked():
            content_3 += self.W_hypo_check.text() + ', '

        if self.W_a_akinesis_check.isChecked():
            content_3 += self.W_akinesis_check.text() + ', '

        if self.W_a_dyskinesis_check.isChecked():
            content_3 += self.W_dysk_check.text()


        self.auto_paragraph('WW',content)
        self.auto_paragraph('w_i',content_1)
        self.auto_paragraph('w_j', content_2)
        self.auto_paragraph('w_k', content_3)

    def comment_word(self):
        self.auto_paragraph('CC',self.comment.toPlainText())

    def basic_information(self):
        self.auto_paragraph('exam', datetime.now().strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日'))
        self.auto_paragraph('birth',self.birthdate_edit.date().toPyDate(\
            ).strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日'))
        self.auto_paragraph('name',self.name_input.text())
        self.auto_paragraph('ID',self.ID_input.text())
        if self.male_input.isChecked():
            self.auto_paragraph('gender','男')
        elif self.female_input.isChecked():
            self.auto_paragraph('gender', '女')

        #calculate the years old
        diff = relativedelta(datetime.now(), self.birthdate_edit.date().toPyDate())
        year_old = str(diff.years) + '歲' + str(diff.months) +'個月'
        self.auto_paragraph('old',year_old)
        self.auto_paragraph('height',self.height_input.text())
        self.auto_paragraph('weight',self.weight_input.text())

    def nessesary_information(self):
        self.auto_word_cell('a_i',self.IVDd_input.text())
        self.auto_word_cell('b_i', self.PWDd_input.text())
        self.auto_word_cell('c_i', self.LVDd_input.text())
        self.auto_word_cell('d_i', self.LVDs_input.text())
        self.auto_word_cell('e_i', self.Ao_input.text())
        self.auto_word_cell('f_i', self.LA_input.text())
        self.auto_word_cell('g_i', self.LVEF_input.text())
        self.auto_word_cell('h_i', self.TAPSE_input.text())
        self.auto_word_cell('i_i', self.EPSS_input.text())
        self.auto_word_cell('j_i', self.IVCexp_input.text())
        self.auto_word_cell('k_i', self.IVCixp_input.text())
        try:
            Res_value = '%.2f' %((float(self.IVCexp_input.text())-float(self.IVCixp_input.text()))*100/float(self.IVCexp_input.text()))
        except:
            Res_value = ''

        self.auto_word_cell('l_i', Res_value)  #Respiratory charge = (exp-isp)/esp  (%)
        self.auto_word_cell('m_i', self.RVD_input.text())
        self.auto_word_cell('n_i', self.RVAd_input.text())
        self.auto_word_cell('o_i', self.RVAD2_input.text())
        self.auto_word_cell('p_i', self.RVFAC_input.text())
        self.auto_word_cell('q_i', self.RAA_input.text())
        self.auto_word_cell('r_i', self.LVEF_input2.text())
        self.auto_word_cell('s_i', self.EA_input.text())
        self.auto_word_cell('t_i', self.MVEv_input.text())
        # self.auto_word_cell('u_i', self..text())
        self.auto_word_cell('v_i', self.Eneed_input.text())
        self.auto_word_cell('w_i', self.Elat_input.text())
        try:
            mean_e_value = '%.2f' %((float(self.Eneed_input.text())+float(self.Elat_input.text()))/2)
        except:
            mean_e_value = ''

        self.auto_word_cell('x_i', mean_e_value)   #mean e'
        try:
            E_e_value = '%.2f' %(float(self.MVEv_input.text())/float(mean_e_value))
        except:
            E_e_value = ''

        self.auto_word_cell('y_i', E_e_value)
        self.auto_word_cell('z_i', self.RVS_input.text())
        self.auto_word_cell('a_j', self.PVSD_combo.currentText())
        self.auto_word_cell('c_j',self.MV_input.text())
        self.auto_word_cell('b_j',self.AsAo_input.text())
        self.auto_word_cell('d_j',self.LVM_index_input.text())
        self.auto_word_cell('z_j',self.MVDt_input.text())
        if self.Per_eff_check.isChecked():
            content = ''
            content += self.Per_combo.currentText()
            if self.Per_combo.currentText()=='nil':
                pass
            else:
                content += '-('+self.Per_eff_combo.currentText()+')' + ' '+self.Per_eff_input.text()+' cm '+ self.with_without_combo.currentText()+' echo-tamponade-sign'
            self.auto_word_cell('p_j',content)

        if self.M_mode_others_check.isChecked():
            self.auto_word_cell('input', self.M_mode_others.toPlainText())
        else:
            self.auto_word_cell('input', '')

    def auto_word_cell(self,replaced_word, content):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'DFKai-SB'
        font.bold=False
        for row in self.table[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replaced_word in paragraph.text:
                        temp = paragraph.text
                        temp = temp.replace(replaced_word,content)
                        paragraph.text = temp
                        paragraph.style = 'Normal'

    def auto_paragraph(self,replaced_word,content):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'DFKai-SB'
        for paragraph in self.document.paragraphs:
            inline = paragraph.runs
            for i in range(len(inline)):
                if replaced_word in inline[i].text:
                    temp = inline[i].text
                    temp = temp.replace(replaced_word, content)
                    inline[i].text = temp


    def reset_input(self):
        self.document = docx.Document('example2.docx')
        self.table = self.document.tables
        self.height_input.setText('')
        self.name_input.setText('')
        self.ID_input.setText('')
        self.weight_input.setText('')
        self.old_input.setText('')
        self.IVDd_input.setText('')
        self.PWDd_input.setText('')
        self.LVDd_input.setText('')
        self.LVDs_input.setText('')
        self.Ao_input.setText('')
        self.LA_input.setText('')
        self.LVEF_input.setText('')
        self.TAPSE_input.setText('')
        self.EPSS_input.setText('')
        self.IVCexp_input.setText('')
        self.IVCixp_input.setText('')
        self.Res_input.setText('')
        self.RVD_input.setText('')
        self.RVAd_input.setText('')
        self.RVAD2_input.setText('')
        self.RVFAC_input.setText('')
        self.RAA_input.setText('')
        self.LVEF_input.setText('')
        self.Per_eff_input.setText('')
        self.MV_input.setText('')
        self.EA_input.setText('')
        self.E_input.setText('')
        self.MVEv_input.setText('')
        self.MVDt_input.setText('')
        self.Eneed_input.setText('')
        self.Elat_input.setText('')
        self.mean_e_input.setText('')
        self.E_e_input.setText('')
        self.RVS_input.setText('')
        self.M_MVA_input.setText('')
        self.M_meantrans_input.setText('')
        self.M_MRarea_input.setText('')
        self.M_MRLA_input.setText('')
        self.M_Veno_input.setText('')
        self.M_ERO_input.setText('')
        self.A_AVA_input.setText('')
        self.A_AVPG_input.setText('')
        self.A_peckPG_input.setText('')
        self.A_AVVmax_input.setText('')
        self.A_jet_input.setText('')
        self.A_PHT_input.setText('')
        self.A_veno_input.setText('')
        self.T_TRarea_input.setText('')
        self.T_TRV_input.setText('')
        self.LVEF_input2.setText('')
        self.M_mode_others_check.setCheckState(0)
        self.M_mode_others.setPlainText('')
        self.Per_eff_check.setCheckState(0)
        self.MV_check.setCheckState(0)
        self.Per_eff_combo.setCurrentIndex(0)
        self.with_without_combo.setCurrentIndex(0)
        self.PVSD_combo.setCurrentIndex(0)
        self.M_normal_check.setCheckState(0)
        self.M_BMV_check.setCheckState(0)
        self.M_MMV_check.setCheckState(0)
        self.M_MMVp_check.setCheckState(0)
        self.M_MMVp_combo.setCurrentIndex(0)
        self.M_MMVp_combo1.setCurrentIndex(0)
        self.M_mitral_check.setCheckState(0)
        self.M_scler_check.setCheckState(0)
        self.M__fibro_check.setCheckState(0)
        self.M_myxo_check.setCheckState(0)
        self.M_chordae_check.setCheckState(0)
        self.M_chordae_check_2.setCurrentIndex(0)
        self.M_MVA_check.setCheckState(0)
        self.M_MR_check.setCheckState(0)
        self.M_MR_combo.setCurrentIndex(0)
        self.M_others_check.setCheckState(0)
        self.M_others.setPlainText('')
        self.A_normal_check.setCheckState(0)
        self.A_BMV_check.setCheckState(0)
        self.A_MMV_check.setCheckState(0)
        self.A_scler_check.setCheckState(0)
        self.A_mitral_check.setCheckState(0)
        self.A_bicu_check.setCheckState(0)
        self.A_tricu_check.setCheckState(0)
        self.A_AVA_check.setCheckState(0)
        self.A_meantrans_check.setCheckState(0)
        self.A_AR_check.setCheckState(0)
        self.A_AR_combo.setCurrentIndex(0)
        self.A_dias_combo.setCurrentIndex(0)
        self.A_others_check.setCheckState(0)
        self.A_others.setPlainText('')
        self.T_normal_check.setCheckState(0)
        self.T_pro_check.setCheckState(0)
        self.T_anter_check.setCheckState(0)
        self.T_post_check.setCheckState(0)
        self.T_septal_check.setCheckState(0)
        self.T_TR_check.setCheckState(0)
        self.T_TR_combo.setCurrentIndex(0)
        self.T_TRV_check.setCheckState(0)
        self.T_TV_check.setCheckState(0)
        self.T_others_check.setCheckState(0)
        self.T_others.setPlainText('')
        self.P_normal_check.setCheckState(0)
        self.P_PR_check.setCheckState(0)
        self.P_others.setPlainText('')
        self.W_normal_check.setCheckState(0)
        self.W_abnormal_check.setCheckState(0)
        self.W_basal_check.setCheckState(0)
        self.W_midcavity_check.setCheckState(0)
        self.W_apical_check.setCheckState(0)
        self.W_anterior_check.setCheckState(0)
        self.W_septal_check.setCheckState(0)
        self.W_inferior_check.setCheckState(0)
        self.W_posterior_check.setCheckState(0)
        self.W_lateral_check.setCheckState(0)
        self.W_hypo_check.setCheckState(0)
        self.W_akinesis_check.setCheckState(0)
        self.W_dysk_check.setCheckState(0)
        self.P_PR_check.setCheckState(0)
        self.comment.setPlainText('')



if __name__=="__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = mywindow()
    window.show()
    sys.exit(app.exec_())